import os
import re
import sqlite3
from PyQt6.QtCore import QObject, pyqtSignal
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from database.db_manager import fetch_questions_by_ids, DB_PATH  # 确保 DB_PATH 在 db_manager.py 中导出

class ExportWorker(QObject):
    """
    后台文档导出 Worker，使用 docxtpl 渲染模板。
    接收题目 id 列表和导出选项，run() 会发出 progress，最后发出 finished。
    """
    progress = pyqtSignal(int, str)
    finished = pyqtSignal()
    preview = pyqtSignal(str)  # 用于实时发送预览内容

    def __init__(self, codes: list[int], opts: dict):
        super().__init__()
        self.codes = codes
        self.opts = opts

    def run(self):
        # 1. 确保输出目录存在
        out_dir = "output"
        os.makedirs(out_dir, exist_ok=True)

        # 2. 读取模板
        tpl_name = self.opts.get("template", "")
        tpl_path = os.path.join("templates", tpl_name)
        self.progress.emit(5, "加载模板…")
        doc = DocxTemplate(tpl_path)

        # 3. 从数据库取题（仅 questions 表字段）
        self.progress.emit(10, "读取题目详情…")
        records = fetch_questions_by_ids(self.codes)

        # —— 新增：清理所有 [IMAGE_n] 占位符，避免在 content_text 中残留 —— 
        for rec in records:
            rec["content_text"] = re.sub(r"\[IMAGE_\d+\]", "", rec.get("content_text", ""))

        # 3.1 额外从 question_images 与 question_formulas 拉数据
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        if not self.codes:
            conn.close()
            return

        placeholders = ",".join("?" for _ in self.codes)

        # 拉图片路径
        cursor.execute(
            f"SELECT question_id, image_path FROM question_images WHERE question_id IN ({placeholders})",
            self.codes
        )
        img_rows = cursor.fetchall()
        img_map = {}
        for qid, path in img_rows:
            # 确保使用正斜杠（替换反斜杠）
            path = path.replace("\\", "/")
            img_map.setdefault(qid, []).append(path)

        # 拉公式（存的是 MathML 内容，若你有公式图片路径也可改成相应字段）
        cursor.execute(
            f"SELECT question_id, content FROM question_formulas WHERE question_id IN ({placeholders})",
            self.codes
        )
        math_rows = cursor.fetchall()
        math_map = {}
        for qid, content in math_rows:
            math_map.setdefault(qid, []).append(content)

        conn.close()

        # 3.2 生成 options、InlineImage 列表
        #    假设你的 media 目录在项目根
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
        media_dir    = os.path.join(project_root, "media")

        for rec in records:
            qid = rec["id"]

            # 选项拼接
            rec["options"] = "\t".join(filter(None, [
                rec.get("option_a", ""),
                rec.get("option_b", ""),
                rec.get("option_c", ""),
                rec.get("option_d", "")
            ]))

            # 图片 InlineImage 列表
            rec["images"] = []
            for rel in img_map.get(qid, []):
                # 兼容绝对与相对路径
                path = rel if os.path.isabs(rel) else os.path.join(media_dir, os.path.normpath(rel))
                if os.path.exists(path):
                    rec["images"].append(InlineImage(doc, path, width=Mm(40)))
                else:
                    self.progress.emit(15, f"⚠️ 图片文件不存在：{path}")

            # 公式：如果内容是 MathML，这里暂时只能把内容原样注入上下文
            # 如果你有公式图片，请同 images 一样处理
            rec["maths"] = math_map.get(qid, [])

            # 实时预览部分 - 发送预览信号
            preview_text = f"题目 {qid}: {rec.get('content_text', '')[:50]}"
            self.preview.emit(preview_text)

        # 4. 按题型分组
        mapping = {
            "单选": "single_choice",
            "多选": "multiple_choice",
            "判断": "judgment",
            "简答": "short_answer",
            "计算": "calculation"
        }
        groups = {v: [] for v in mapping.values()}
        for rec in records:
            grp = mapping.get(rec.get("question_type", ""), "")
            if grp:
                groups[grp].append(rec)

        # 5. 准备渲染上下文
        context = {
            "job_name":         self.opts.get("job_name", ""),
            "level_name":       self.opts.get("level_name", ""),
            "header":           self.opts.get("header", ""),
            "with_ans":         self.opts.get("with_ans", False),
            "show_recognition": self.opts.get("new_page", False),
            **groups
        }

        # 6. 渲染
        self.progress.emit(50, "渲染文档…")
        doc.render(context)

        # 7. 动态文件名（防重名）
        job   = context["job_name"] or "job"
        level = context["level_name"] or "level"
        filename = f"{job}_{level}"
        if context.get("show_recognition"):
            filename += "_认定点"
        if context.get("with_ans"):
            filename += "_答案解析"

        base, ext = filename, ".docx"
        out_path = os.path.join(out_dir, base + ext)
        counter = 1
        while os.path.exists(out_path):
            out_path = os.path.join(out_dir, f"{base}_{counter}{ext}")
            counter += 1

        # 8. 保存初版
        doc.save(out_path)

        # 9. 清理空段落并合并 [T] 行
        self.progress.emit(80, "清理空段落并合并题头…")
        try:
            doc2 = Document(out_path)

            # 删除全空段落
            for para in list(doc2.paragraphs):
                if not para.text.strip():
                    para._element.getparent().remove(para._element)

            # 合并以 [T] 结尾的段落与下一段
            paras = doc2.paragraphs
            i = 0
            while i < len(paras) - 1:
                p1 = paras[i]
                if p1.text.strip().endswith("[T]"):
                    p2 = paras[i + 1]
                    for run in p2.runs:
                        new_run = p1.add_run(run.text)
                        new_run.bold      = run.bold
                        new_run.italic    = run.italic
                        new_run.underline = run.underline
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                    p2._element.getparent().remove(p2._element)
                    paras = doc2.paragraphs
                else:
                    i += 1

            doc2.save(out_path)
        except Exception as e:
            self.progress.emit(80, f"⚠️ 清理/合并异常：{e}")

        # 10. 完成
        self.progress.emit(100, f"导出完成：{os.path.basename(out_path)}")
        self.finished.emit()
